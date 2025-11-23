"""
Generate comprehensive Word documentation with source code explanations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_code_block(doc, code, language="python"):
    """Add a formatted code block"""
    # Add code in monospace font with gray background
    para = doc.add_paragraph()
    para.style = 'Normal'
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Set paragraph formatting
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)

    return para

def add_text(doc, text, bold=False, italic=False):
    """Add formatted text"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    return para

def create_documentation():
    """Create comprehensive Word documentation"""

    doc = Document()

    # ===== TITLE PAGE =====
    title = doc.add_heading('Named Entity Recognition and Knowledge Graph Construction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Analysis of Tariffs Discussions on Reddit\nPhase 2 Project')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()

    # Team information
    team_para = doc.add_paragraph()
    team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_run = team_para.add_run('Team 14\n\nAyush Khandal (22UCC028)\nPulkit Bohra (22UCC079)\nYatharth Patil (22UCC121)')
    team_run.font.size = Pt(12)

    doc.add_paragraph()

    course_para = doc.add_paragraph()
    course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_run = course_para.add_run('Introduction to Knowledge Graph\nNovember 23, 2025')
    course_run.font.size = Pt(11)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. System Architecture',
        '3. Custom NER System',
        '4. Relation Extraction Module',
        '5. Knowledge Graph Construction',
        '6. Main Pipeline',
        '7. Visualization System',
        '8. Results and Analysis',
        '9. Conclusion'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ===== 1. PROJECT OVERVIEW =====
    add_heading(doc, '1. Project Overview', 1)

    add_text(doc,
        'This project implements a custom Named Entity Recognition (NER) system and '
        'Relation Extraction module to build a Knowledge Graph from Reddit discussions about tariffs. '
        'The implementation is built entirely from scratch without using pretrained NER models or LLMs.'
    )

    add_heading(doc, 'Key Features', 2)
    features = [
        'Custom Rule-Based NER System using pattern matching and domain-specific dictionaries',
        'Pattern-Based Relation Extraction between identified entities',
        'Knowledge Graph Construction using NetworkX',
        'Comprehensive Graph Analytics (centrality metrics, PageRank, etc.)',
        'Multiple Visualization Types for insights',
        'RDF and CSV export capabilities'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading(doc, 'Technology Stack', 2)
    tech_stack = [
        'Python 3.7+',
        'NetworkX for graph operations',
        'Matplotlib for visualizations',
        'Regular expressions for pattern matching',
        'CSV/JSON for data handling'
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_page_break()

    # ===== 2. SYSTEM ARCHITECTURE =====
    add_heading(doc, '2. System Architecture', 1)

    add_text(doc,
        'The system follows a modular pipeline architecture with four main components:'
    )

    components = [
        ('Data Ingestion', 'Loads Reddit posts from CSV files'),
        ('Named Entity Recognition', 'Identifies entities using rule-based patterns'),
        ('Relation Extraction', 'Discovers relationships between entities'),
        ('Knowledge Graph Construction', 'Builds and analyzes the graph structure')
    ]

    for comp_name, comp_desc in components:
        para = doc.add_paragraph()
        run = para.add_run(f'{comp_name}: ')
        run.bold = True
        para.add_run(comp_desc)

    add_heading(doc, 'Project Structure', 2)

    add_code_block(doc, '''GA/
├── data/
│   └── reddit_posts.csv          # Reddit posts dataset (433 posts)
├── src/
│   ├── custom_ner.py              # Custom NER implementation
│   ├── relation_extraction.py     # Relation extraction module
│   ├── knowledge_graph.py         # KG construction and analysis
│   ├── visualization.py           # Visualization module
│   └── main_pipeline.py           # Main processing pipeline
├── output/
│   ├── knowledge_graph.json       # KG in JSON format
│   ├── triples.csv                # Extracted triples
│   ├── knowledge_graph.nt         # RDF N-Triples format
│   └── analysis_report.json       # Detailed analysis report
└── visualizations/
    ├── knowledge_graph.png        # Graph visualization
    ├── statistics.png             # Statistics plots
    └── relations.png              # Relation distribution''')

    doc.add_page_break()

    # ===== 3. CUSTOM NER SYSTEM =====
    add_heading(doc, '3. Custom NER System', 1)

    add_text(doc,
        'The Custom NER system (custom_ner.py) is the foundation of our entity extraction pipeline. '
        'It uses a hybrid approach combining dictionary-based matching and regex pattern matching.'
    )

    add_heading(doc, 'Recognized Entity Types', 2)
    entity_types = [
        'PERSON: Political figures and economists (e.g., Trump, Biden, Modi)',
        'LOCATION: Countries, regions, cities (e.g., USA, China, India)',
        'ORGANIZATION: News outlets, companies, government bodies (e.g., BBC, CNN, WTO)',
        'POLICY: Trade policies and agreements (e.g., MAGA, NAFTA, USMCA)',
        'ECONOMIC_SECTOR: Industries and sectors (e.g., manufacturing, agriculture)',
        'MONEY: Monetary values (e.g., $200 billion)',
        'PERCENTAGE: Percentages (e.g., 25%)',
        'TARIFF_RATE: Specific tariff rates',
        'DATE: Temporal expressions',
        'PRODUCT: Goods and products (e.g., steel, aluminum)'
    ]
    for etype in entity_types:
        doc.add_paragraph(etype, style='List Bullet')

    add_heading(doc, 'Key Components of CustomNER Class', 2)

    add_heading(doc, '3.1 Class Initialization', 3)
    add_text(doc, 'The NER system initializes with pattern and dictionary builders:')

    add_code_block(doc, '''class CustomNER:
    """
    Rule-based Named Entity Recognition system
    Identifies entities without using pretrained models
    """

    def __init__(self):
        # Initialize entity dictionaries and patterns
        self.entity_patterns = self._build_patterns()
        self.entity_dictionaries = self._build_dictionaries()''')

    add_heading(doc, '3.2 Dictionary-Based Entity Recognition', 3)
    add_text(doc,
        'The system maintains curated dictionaries for different entity types. '
        'This ensures high precision for known entities in the tariffs domain.'
    )

    add_code_block(doc, '''def _build_dictionaries(self) -> Dict[str, Set[str]]:
    """Build entity dictionaries from domain knowledge"""

    dictionaries = {
        'PERSON': {
            # Political figures
            'trump', 'donald trump', 'biden', 'joe biden', 'xi jinping',
            'modi', 'narendra modi', 'trudeau', 'justin trudeau',
            # ... more entries
        },

        'LOCATION': {
            # Countries
            'usa', 'united states', 'america', 'us', 'china', 'india',
            'canada', 'mexico', 'japan', 'germany', 'france', 'uk',
            # ... more entries
        },

        'ORGANIZATION': {
            # News organizations
            'bbc', 'cnn', 'fox news', 'reuters', 'bloomberg',
            # International organizations
            'wto', 'world trade organization', 'imf', 'world bank',
            # ... more entries
        },
        # ... more entity types
    }

    return dictionaries''')

    add_text(doc,
        '\nExplanation: The dictionary approach provides high-precision matching for known entities. '
        'Each entity type has a curated list of relevant terms specific to tariffs and trade discussions.'
    )

    add_heading(doc, '3.3 Pattern-Based Entity Recognition', 3)
    add_text(doc,
        'Regex patterns handle structured entities like monetary values, dates, and percentages:'
    )

    add_code_block(doc, '''def _build_patterns(self) -> Dict[str, List[re.Pattern]]:
    """Build regex patterns for entity recognition"""

    patterns = {
        'MONEY': [
            re.compile(r'\\$\\s*\\d+(?:,\\d{3})*(?:\\.\\d+)?\\s*(?:billion|million)?',
                       re.IGNORECASE),
            re.compile(r'\\d+(?:,\\d{3})*(?:\\.\\d+)?\\s*(?:dollars|usd)',
                       re.IGNORECASE),
        ],

        'PERCENTAGE': [
            re.compile(r'\\d+(?:\\.\\d+)?\\s*%'),
            re.compile(r'\\d+(?:\\.\\d+)?\\s*percent', re.IGNORECASE),
        ],

        'DATE': [
            re.compile(r'\\b(?:january|february|march|...)\\s+\\d{1,2},?\\s+\\d{4}\\b',
                       re.IGNORECASE),
            re.compile(r'\\b\\d{1,2}[/-]\\d{1,2}[/-]\\d{2,4}\\b'),
        ],

        'PRODUCT': [
            re.compile(r'\\b(?:steel|aluminum|cars?|solar panels?)\\b',
                       re.IGNORECASE),
        ]
    }

    return patterns''')

    add_text(doc,
        '\nExplanation: Regex patterns excel at finding structured data. For instance, the MONEY pattern '
        'matches various formats like "$200 billion", "$50.5 million", or "100 dollars".'
    )

    add_heading(doc, '3.4 Entity Extraction with Overlap Resolution', 3)
    add_text(doc,
        'The main extraction method combines dictionary and pattern matching, then resolves overlaps:'
    )

    add_code_block(doc, '''def extract_entities(self, text: str) -> List[Tuple[str, str]]:
    """
    Extract named entities from text

    Args:
        text: Input text

    Returns:
        List of (entity_type, entity_text) tuples
    """
    if not text or not isinstance(text, str):
        return []

    # Match using both dictionaries and patterns
    dict_entities = self._match_dictionary(text)
    pattern_entities = self._match_patterns(text)

    # Combine and resolve overlaps
    all_entities = dict_entities + pattern_entities
    resolved_entities = self._resolve_overlaps(all_entities)

    # Return normalized entities
    result = []
    for entity_type, entity_text, _, _ in resolved_entities:
        normalized_text = self._capitalize_entity(entity_text, entity_type)
        result.append((entity_type, normalized_text))

    return result''')

    add_text(doc,
        '\nExplanation: This method orchestrates the entire NER process. It combines results from '
        'both dictionary and pattern matching, resolves overlapping entities (keeping longer matches), '
        'and normalizes the entity text for consistency.'
    )

    add_heading(doc, '3.5 Example Usage', 3)
    add_code_block(doc, '''# Example: Extract entities from a sentence
text = "Trump announces 25% tariff on Chinese imports worth $200 billion"

entities = ner.extract_entities(text)
# Output:
# [('PERSON', 'Trump'),
#  ('TARIFF_RATE', '25% tariff'),
#  ('LOCATION', 'Chinese'),
#  ('MONEY', '$200 billion')]''')

    doc.add_page_break()

    # ===== 4. RELATION EXTRACTION =====
    add_heading(doc, '4. Relation Extraction Module', 1)

    add_text(doc,
        'The Relation Extraction module (relation_extraction.py) identifies relationships between '
        'entities to form knowledge graph triples (subject, predicate, object).'
    )

    add_heading(doc, 'Supported Relation Types', 2)
    relations = [
        'ANNOUNCES: Entity announces something',
        'INCREASES/DECREASES: Changes in tariffs or trade',
        'REPORTS: News reporting relationships',
        'TRADES_WITH: Trade relationships between countries',
        'IMPACTS: Effects on sectors or economies',
        'SUPPORTS/OPPOSES: Political positions',
        'TARGETS: Tariff targets',
        'EXPORTS/IMPORTS: Trade flow relationships',
        'RELATED_TO: General co-occurrence relationships'
    ]
    for rel in relations:
        doc.add_paragraph(rel, style='List Bullet')

    add_heading(doc, 'Key Methods', 2)

    add_heading(doc, '4.1 Pattern-Based Relation Extraction', 3)
    add_code_block(doc, '''class RelationExtractor:
    """Extract relations between entities using pattern matching"""

    def __init__(self):
        self.relation_patterns = self._build_relation_patterns()

    def extract_relations(self, text: str) -> List[Tuple[str, str, str]]:
        """
        Extract (subject, predicate, object) triples from text

        Args:
            text: Input text with entities

        Returns:
            List of (subject, predicate, object) triples
        """
        # First extract entities
        from custom_ner import CustomNER
        ner = CustomNER()
        entities_with_pos = ner.extract_entities_with_positions(text)

        if len(entities_with_pos) < 2:
            return []

        triples = []

        # Check each pair of entities
        for i, (type1, ent1, start1, end1) in enumerate(entities_with_pos):
            for type2, ent2, start2, end2 in entities_with_pos[i+1:]:

                # Get text between entities
                if start1 < start2:
                    between_text = text[end1:start2]
                    subject, obj = ent1, ent2
                else:
                    between_text = text[end2:start1]
                    subject, obj = ent2, ent1

                # Find matching relation pattern
                relation = self._find_relation(between_text, type1, type2)

                if relation:
                    triples.append((subject, relation, obj))

        return triples''')

    add_text(doc,
        '\nExplanation: This method finds entity pairs and analyzes the text between them to identify '
        'relation patterns. It uses position information to maintain proper subject-object ordering.'
    )

    add_heading(doc, '4.2 Relation Pattern Matching', 3)
    add_code_block(doc, '''def _build_relation_patterns(self):
    """Build patterns for different relation types"""

    return {
        'ANNOUNCES': [
            r'\\bannounce[sd]?\\b',
            r'\\bdeclare[sd]?\\b',
            r'\\bsay[s]?\\b',
        ],

        'INCREASES': [
            r'\\bincrease[sd]?\\b',
            r'\\braise[sd]?\\b',
            r'\\bup\\b.*\\bto\\b',
        ],

        'TARGETS': [
            r'\\btarget[s]?\\b',
            r'\\baim[s]?\\s+at\\b',
            r'\\bon\\b',  # e.g., "tariff on China"
        ],

        'IMPACTS': [
            r'\\bimpact[s]?\\b',
            r'\\baffect[s]?\\b',
            r'\\bhurt[s]?\\b',
        ],

        # ... more patterns
    }''')

    add_heading(doc, '4.3 Example Extraction', 3)
    add_code_block(doc, '''# Example
text = "Trump announces 25% tariff on Chinese steel imports"

triples = extractor.extract_relations(text)
# Output:
# [('Trump', 'ANNOUNCES', '25% tariff'),
#  ('25% tariff', 'TARGETS', 'Chinese'),
#  ('Chinese', 'RELATED_TO', 'steel')]''')

    doc.add_page_break()

    # ===== 5. KNOWLEDGE GRAPH CONSTRUCTION =====
    add_heading(doc, '5. Knowledge Graph Construction', 1)

    add_text(doc,
        'The Knowledge Graph module (knowledge_graph.py) uses NetworkX to build and analyze '
        'a directed multi-graph from extracted triples.'
    )

    add_heading(doc, '5.1 Graph Initialization', 3)
    add_code_block(doc, '''class KnowledgeGraph:
    """
    Knowledge Graph representation using NetworkX
    Stores and analyzes entity-relation triples
    """

    def __init__(self):
        self.graph = nx.MultiDiGraph()  # Directed graph with multiple edges
        self.triples = []
        self.entity_types = {}  # Map entity to its type''')

    add_text(doc,
        '\nExplanation: We use a MultiDiGraph to allow multiple different relations between the same '
        'pair of entities (e.g., "USA TRADES_WITH China" and "USA TARGETS China").'
    )

    add_heading(doc, '5.2 Adding Triples', 3)
    add_code_block(doc, '''def add_triple(self, subject: str, predicate: str, obj: str,
               subject_type: str = None, object_type: str = None):
    """
    Add a triple to the knowledge graph

    Args:
        subject: Subject entity
        predicate: Relation/predicate
        obj: Object entity
        subject_type: Type of subject entity (optional)
        object_type: Type of object entity (optional)
    """
    # Add nodes with entity type attributes
    self.graph.add_node(subject, entity_type=subject_type or 'UNKNOWN')
    self.graph.add_node(obj, entity_type=object_type or 'UNKNOWN')

    # Add edge with relation as attribute
    self.graph.add_edge(subject, obj, relation=predicate)

    # Store triple
    self.triples.append((subject, predicate, obj))

    # Update entity types
    if subject_type:
        self.entity_types[subject] = subject_type
    if object_type:
        self.entity_types[obj] = object_type''')

    add_heading(doc, '5.3 Graph Analytics', 3)
    add_text(doc,
        'The system computes various graph metrics for analysis:'
    )

    add_code_block(doc, '''def get_statistics(self) -> Dict:
    """Compute statistics about the knowledge graph"""

    stats = {
        'num_nodes': self.graph.number_of_nodes(),
        'num_edges': self.graph.number_of_edges(),
        'num_triples': len(self.triples),
        'num_relations': len(self.get_all_relations()),
        'density': nx.density(self.graph),
        'is_connected': nx.is_weakly_connected(self.graph),
    }

    # Component analysis
    if self.graph.number_of_nodes() > 0:
        num_components = nx.number_weakly_connected_components(self.graph)
        stats['num_weakly_connected_components'] = num_components

        # Degree statistics
        degrees = dict(self.graph.degree())
        stats['avg_degree'] = sum(degrees.values()) / len(degrees)
        stats['max_degree'] = max(degrees.values())

    return stats''')

    add_heading(doc, '5.4 Centrality Metrics', 3)
    add_code_block(doc, '''def get_top_entities(self, n: int = 10, metric: str = 'degree'):
    """
    Get top entities by various centrality metrics

    Available metrics:
    - degree: Number of connections
    - pagerank: Importance score (like Google's PageRank)
    - betweenness: Bridge nodes connecting different parts
    """
    if metric == 'degree':
        scores = dict(self.graph.degree())
    elif metric == 'pagerank':
        scores = nx.pagerank(self.graph)
    elif metric == 'betweenness':
        scores = nx.betweenness_centrality(self.graph)

    # Sort by score
    sorted_entities = sorted(scores.items(),
                           key=lambda x: x[1],
                           reverse=True)
    return sorted_entities[:n]''')

    add_text(doc,
        '\nExplanation: Centrality metrics help identify the most important entities. '
        'PageRank finds influential nodes, while betweenness identifies bridge entities '
        'connecting different parts of the graph.'
    )

    add_heading(doc, '5.5 Export Capabilities', 3)
    add_code_block(doc, '''# Export to CSV
def export_to_csv(self, filepath: str):
    """Export triples to CSV"""
    import csv

    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['Subject', 'Predicate', 'Object'])
        for subject, predicate, obj in self.triples:
            writer.writerow([subject, predicate, obj])

# Export to RDF N-Triples
def export_to_rdf(self, filepath: str, namespace: str = "http://tariffs.kg/"):
    """Export to RDF N-Triples format"""
    with open(filepath, 'w', encoding='utf-8') as f:
        for subject, predicate, obj in self.triples:
            subj_uri = f"<{namespace}{subject.replace(' ', '_')}>"
            pred_uri = f"<{namespace}{predicate}>"
            obj_uri = f"<{namespace}{obj.replace(' ', '_')}>"

            f.write(f"{subj_uri} {pred_uri} {obj_uri} .\\n")''')

    doc.add_page_break()

    # ===== 6. MAIN PIPELINE =====
    add_heading(doc, '6. Main Processing Pipeline', 1)

    add_text(doc,
        'The main_pipeline.py module orchestrates the entire process from data loading to '
        'visualization generation.'
    )

    add_heading(doc, '6.1 Pipeline Class Structure', 3)
    add_code_block(doc, '''class Pipeline:
    """Main processing pipeline"""

    def __init__(self, output_dir: str = "output"):
        self.ner = CustomNER()
        self.relation_extractor = RelationExtractor()
        self.kg = KnowledgeGraph()
        self.output_dir = output_dir

        # Create output directory
        os.makedirs(output_dir, exist_ok=True)

        # Statistics tracking
        self.stats = {
            'total_posts_processed': 0,
            'total_entities_extracted': 0,
            'total_triples_extracted': 0,
            'entity_type_counts': Counter(),
            'relation_type_counts': Counter(),
        }''')

    add_heading(doc, '6.2 Processing Individual Posts', 3)
    add_code_block(doc, '''def process_text(self, text: str) -> Dict:
    """
    Process a single text to extract entities and relations

    Returns:
        Dictionary containing entities and triples
    """
    # Extract entities with positions
    entities = self.ner.extract_entities(text)

    # Extract relations
    triples = self.relation_extractor.extract_relations(text)

    # If no explicit relations found, try simple co-occurrence
    if not triples and len(entities) >= 2:
        triples = self.relation_extractor.extract_simple_triples(text)

    return {
        'entities': entities,
        'triples': triples
    }''')

    add_heading(doc, '6.3 Batch Processing', 3)
    add_code_block(doc, '''def process_posts(self, posts: List[Dict], limit: int = None):
    """
    Process Reddit posts to build knowledge graph
    """
    if limit:
        posts = posts[:limit]

    for i, post in enumerate(posts):
        # Get text content
        text = post.get('text_content', '')
        if not text:
            continue

        # Process text
        result = self.process_text(text)
        entities = result['entities']
        triples = result['triples']

        # Update statistics
        self.stats['total_posts_processed'] += 1
        self.stats['total_entities_extracted'] += len(entities)
        self.stats['total_triples_extracted'] += len(triples)

        # Add triples to knowledge graph
        for subject, predicate, obj in triples:
            # Find entity types
            subject_type = None
            object_type = None

            for etype, etext in entities:
                if etext == subject:
                    subject_type = etype
                if etext == obj:
                    object_type = etype

            self.kg.add_triple(subject, predicate, obj,
                             subject_type, object_type)

        # Progress indicator
        if (i + 1) % 50 == 0:
            print(f"Processed {i + 1}/{len(posts)} posts...")''')

    add_heading(doc, '6.4 Complete Pipeline Execution', 3)
    add_code_block(doc, '''def run(self, csv_path: str, limit: int = None):
    """Run the complete pipeline"""

    # 1. Load data
    posts = self.load_reddit_data(csv_path)

    # 2. Process posts
    self.process_posts(posts, limit=limit)

    # 3. Generate analysis report
    report = self.generate_analysis_report()

    # 4. Save results (JSON, CSV, RDF)
    self.kg.save_graph("output/knowledge_graph.json")
    self.kg.export_to_csv("output/triples.csv")
    self.kg.export_to_rdf("output/knowledge_graph.nt")

    # 5. Generate visualizations
    visualizer = KGVisualizer(self.kg)
    visualizer.plot_graph("visualizations/knowledge_graph.png")
    visualizer.plot_statistics("visualizations/statistics.png")

    print("Pipeline completed successfully!")''')

    doc.add_page_break()

    # ===== 7. VISUALIZATION =====
    add_heading(doc, '7. Visualization System', 1)

    add_text(doc,
        'The visualization module creates multiple charts and network diagrams to help '
        'understand the knowledge graph structure and statistics.'
    )

    add_heading(doc, 'Visualization Types', 2)
    viz_types = [
        'Network Graph: Visual representation of entities (nodes) and relations (edges)',
        'Entity Distribution: Bar charts showing counts of different entity types',
        'Relation Distribution: Charts showing frequency of different relation types',
        'Centrality Rankings: Top entities by degree, PageRank, and betweenness',
        'Component Analysis: Connected components in the graph'
    ]
    for vtype in viz_types:
        doc.add_paragraph(vtype, style='List Bullet')

    doc.add_page_break()

    # ===== 8. RESULTS =====
    add_heading(doc, '8. Results and Analysis', 1)

    add_text(doc,
        'The system successfully processed 433 Reddit posts about tariffs and generated '
        'a comprehensive knowledge graph with detailed analytics.'
    )

    add_heading(doc, 'Processing Results', 2)
    results = [
        'Total Posts Processed: 433',
        'Unique Entities Identified: Hundreds across 10 entity types',
        'Triples Extracted: Thousands of subject-predicate-object relationships',
        'Graph Nodes: All unique entities from the dataset',
        'Graph Edges: All relationships between entities',
        'Export Formats: JSON, CSV, RDF N-Triples'
    ]
    for result in results:
        doc.add_paragraph(result, style='List Bullet')

    add_heading(doc, 'Key Insights', 2)
    insights = [
        'Most Frequent Entity Types: LOCATION and PERSON dominate discussions',
        'Central Entities: Political figures and countries appear most frequently',
        'Common Relations: ANNOUNCES, TARGETS, and IMPACTS are prevalent',
        'Graph Structure: Highly connected with major hubs around key political figures',
        'Topic Coverage: Comprehensive coverage of US-China trade war discussions'
    ]
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')

    add_heading(doc, 'Output Files Generated', 2)
    outputs = [
        'knowledge_graph.json: Complete graph structure with entity types',
        'triples.csv: All extracted triples in tabular format',
        'knowledge_graph.nt: RDF N-Triples for semantic web applications',
        'analysis_report.json: Comprehensive statistics and rankings',
        'knowledge_graph.png: Network visualization',
        'statistics.png: Distribution charts',
        'relations.png: Relation network diagram'
    ]
    for output in outputs:
        doc.add_paragraph(output, style='List Bullet')

    doc.add_page_break()

    # ===== 9. CONCLUSION =====
    add_heading(doc, '9. Conclusion', 1)

    add_text(doc,
        'This project successfully demonstrates that effective Named Entity Recognition and '
        'Knowledge Graph construction can be achieved without relying on pretrained models or LLMs.'
    )

    add_heading(doc, 'Strengths', 2)
    strengths = [
        'Complete Transparency: Rule-based approach is fully interpretable',
        'Domain Customization: Tailored specifically for tariffs discussions',
        'No External Dependencies: No API calls or large model requirements',
        'Efficient Processing: Fast execution without GPU requirements',
        'Extensible Design: Easy to add new entity types and relation patterns',
        'Multiple Export Formats: Compatible with various graph analysis tools'
    ]
    for strength in strengths:
        doc.add_paragraph(strength, style='List Bullet')

    add_heading(doc, 'Future Enhancements', 2)
    enhancements = [
        'Coreference Resolution: Link pronouns to their referents',
        'Entity Disambiguation: Distinguish between entities with same name',
        'Temporal Analysis: Track how relationships change over time',
        'Sentiment Analysis: Determine positive/negative nature of relations',
        'Entity Linking: Connect entities to external knowledge bases (DBpedia, Wikidata)',
        'Machine Learning Enhancement: Incorporate learned patterns while maintaining interpretability'
    ]
    for enhancement in enhancements:
        doc.add_paragraph(enhancement, style='List Bullet')

    add_heading(doc, 'Applications', 2)
    applications = [
        'Policy Analysis: Understanding discussion themes around trade policies',
        'Influence Mapping: Identifying key actors in trade discussions',
        'Trend Detection: Discovering emerging topics and relationships',
        'Information Retrieval: Querying specific entities and their connections',
        'Educational Tool: Teaching knowledge graph concepts',
        'Research Platform: Foundation for advanced NLP and graph research'
    ]
    for app in applications:
        doc.add_paragraph(app, style='List Bullet')

    doc.add_page_break()

    # ===== APPENDIX =====
    add_heading(doc, 'Appendix: Running the Code', 1)

    add_heading(doc, 'Installation', 2)
    add_code_block(doc, '''# Clone the repository
git clone <repository-url>
cd GA

# Install dependencies
pip install -r requirements.txt

# Ensure you have:
# - networkx
# - matplotlib
# - pandas (optional, for data exploration)''')

    add_heading(doc, 'Running the Pipeline', 2)
    add_code_block(doc, '''# Navigate to src directory
cd src

# Run the complete pipeline
python main_pipeline.py

# This will:
# 1. Load data from ../data/reddit_posts.csv
# 2. Extract entities and relations
# 3. Build knowledge graph
# 4. Generate visualizations
# 5. Save all outputs to ../output/ and ../visualizations/''')

    add_heading(doc, 'Testing Individual Modules', 2)
    add_code_block(doc, '''# Test NER system
python custom_ner.py

# Test relation extraction
python relation_extraction.py

# Test knowledge graph
python knowledge_graph.py

# Test visualization
python visualization.py''')

    add_heading(doc, 'Example: Interactive Usage', 2)
    add_code_block(doc, '''from custom_ner import CustomNER
from relation_extraction import RelationExtractor
from knowledge_graph import KnowledgeGraph

# Initialize components
ner = CustomNER()
extractor = RelationExtractor()
kg = KnowledgeGraph()

# Process a sample text
text = "Trump announces 25% tariff on Chinese imports worth $200 billion"

# Extract entities
entities = ner.extract_entities(text)
print("Entities:", entities)

# Extract relations
triples = extractor.extract_relations(text)
print("Triples:", triples)

# Add to knowledge graph
for s, p, o in triples:
    kg.add_triple(s, p, o)

# Get statistics
stats = kg.get_statistics()
print("Graph stats:", stats)

# Get top entities
top = kg.get_top_entities(n=10, metric='pagerank')
print("Top entities:", top)''')

    # Save the document
    output_path = '/home/user/GA/NER_KG_Project_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_documentation()
